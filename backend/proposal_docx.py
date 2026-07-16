"""DOCX serializer — the proposal as a real Word document.

Replaces the retired deck/PPTX path. The pricing section is a genuine Word table
(not a text blob), and every section is selectable, editable text — the point of
going DOCX-first.

Also exposes `html_to_blocks`, the small HTML walker shared with the PDF
serializer, so both formats render the same section content identically.
"""

import io
from html.parser import HTMLParser
from typing import Any, Dict, List, Tuple

INK = "141414"
MUTED = "6B6B6B"
LINE = "E5E3DF"


# ----------------------------- Shared HTML → blocks ----------------------------
class _Block:
    __slots__ = ("kind", "runs")

    def __init__(self, kind: str):
        self.kind = kind                    # "para" | "bullet" | "heading"
        self.runs: List[Tuple[str, bool, bool]] = []   # (text, bold, italic)


class _BlockParser(HTMLParser):
    """Turns the proposal's simple HTML (p, ul/ol/li, strong/b, em/i, br) into a flat
    list of blocks with inline bold/italic — the common shape both serializers need."""

    def __init__(self):
        super().__init__()
        self.blocks: List[_Block] = []
        self._bold = 0
        self._italic = 0
        self._cur: _Block = None

    def _open(self, kind: str):
        self._cur = _Block(kind)
        self.blocks.append(self._cur)

    def handle_starttag(self, tag, attrs):
        if tag in ("p",):
            self._open("para")
        elif tag == "li":
            self._open("bullet")
        elif tag in ("strong", "b"):
            self._bold += 1
        elif tag in ("em", "i"):
            self._italic += 1
        elif tag == "br":
            if self._cur:
                self._cur.runs.append(("\n", False, False))

    def handle_endtag(self, tag):
        if tag in ("strong", "b"):
            self._bold = max(0, self._bold - 1)
        elif tag in ("em", "i"):
            self._italic = max(0, self._italic - 1)
        elif tag in ("p", "li"):
            self._cur = None

    def handle_data(self, data):
        if not data.strip() and data != " ":
            return
        if self._cur is None:
            self._open("para")
        self._cur.runs.append((data, self._bold > 0, self._italic > 0))


def html_to_blocks(html: str) -> List[_Block]:
    p = _BlockParser()
    p.feed(html or "")
    return [b for b in p.blocks if b.runs]


# ----------------------------- DOCX -------------------------------------------
def build_docx(proposal: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title block
    title = doc.add_paragraph()
    run = title.add_run(proposal.get("topic", "Proposal"))
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(INK)
    client = (proposal.get("client_facts") or {}).get("company")
    if client:
        sub = doc.add_paragraph()
        r = sub.add_run(f"Prepared for {client}")
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()

    for section in proposal.get("sections", []):
        h = doc.add_heading(section.get("heading", ""), level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor.from_string(INK)

        if section.get("slot") == "pricing_table":
            _add_pricing_table(doc, proposal.get("pricing") or {})
            note = section.get("html", "")
            for blk in html_to_blocks(note):
                _add_block(doc, blk)
        else:
            blocks = html_to_blocks(section.get("html", ""))
            if not blocks:
                p = doc.add_paragraph()
                r = p.add_run("[needs input]")
                r.italic = True
                r.font.color.rgb = RGBColor.from_string(MUTED)
            for blk in blocks:
                _add_block(doc, blk)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_block(doc, blk: _Block):
    from docx.shared import Pt
    if blk.kind == "bullet":
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
    for text, bold, italic in blk.runs:
        if text == "\n":
            p.add_run().add_break()
            continue
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic


def _add_pricing_table(doc, pricing: Dict[str, Any]):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    line_items = pricing.get("line_items") or []
    if not line_items:
        p = doc.add_paragraph()
        r = p.add_run("[needs input: add items to your pricing catalog]")
        r.italic = True
        return

    cur = pricing.get("currency", "USD")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(("Item", "Qty", "Unit price", "Amount")):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True

    for li in line_items:
        cells = table.add_row().cells
        cells[0].text = li["name"] + (f" — {li['description']}" if li.get("description") else "")
        cells[1].text = str(li["qty"])
        cells[2].text = _money(li["unit_price"], cur)
        cells[3].text = _money(li["line_total"], cur)
        for i in (1, 2, 3):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Totals
    def total_row(label, value, bold=False):
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        cells[2].text = ""
        lab = cells[2].paragraphs[0].add_run(label)
        lab.bold = bold
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        val = cells[3].paragraphs[0].add_run(_money(value, cur))
        val.bold = bold
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    total_row("Subtotal", pricing.get("subtotal", 0))
    if pricing.get("discount"):
        total_row(f"Discount ({pricing.get('discount_pct', 0):g}%)", -pricing["discount"])
    total_row("Total", pricing.get("total", 0), bold=True)


def _money(v: Any, currency: str = "USD") -> str:
    sym = {"USD": "$", "EUR": "€", "GBP": "£", "INR": "₹"}.get(currency, "")
    try:
        n = float(v)
    except (TypeError, ValueError):
        return str(v)
    body = f"{abs(n):,.2f}"
    s = f"{sym}{body}" if sym else f"{body} {currency}"
    return f"-{s}" if n < 0 else s
